import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def add_title_and_section1(doc):
    # Add Title
    title = doc.add_paragraph("BÁO CÁO NGHIÊN CỨU KIẾN TRÚC & GIẢI THUẬT HỆ THỐNG", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Dự án: AI-Driven Dynamic Itinerary Optimizer")
    subtitle.paragraph_format.space_after = Pt(36)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1
    doc.add_heading("1. Nghiên cứu Khối Mobile App (React Native & Expo)", level=1)
    
    doc.add_heading("1.1 Xác định các modules chính", level=2)
    p1 = doc.add_paragraph(
        "Phân hệ Mobile App đóng vai trò là Layer 5 trong kiến trúc 5 lớp của hệ thống. "
        "Ứng dụng được xây dựng trên nền tảng React Native sử dụng Expo CLI và Ignite CLI Template, "
        "kết hợp ngôn ngữ TypeScript để đảm bảo tính an toàn dữ liệu. Các module cốt lõi của phân hệ bao gồm:"
    )
    
    doc.add_paragraph("• Module Điều hướng (Navigation System): Sử dụng Expo Router quản lý định tuyến dựa trên cấu trúc thư mục file-system, tăng tốc khả năng mở rộng trang.", style='List Bullet')
    doc.add_paragraph("• Module Bản đồ & Địa lý (Map & Geolocation): Tích hợp thư viện bản đồ để dựng bản đồ vector trực quan, hiển thị lộ trình và POI (Points of Interest).", style='List Bullet')
    doc.add_paragraph("• Module Quản lý trạng thái (State Management): Lưu trữ dữ liệu chuyến đi, thông số cài đặt cá nhân hóa của khách du lịch và quản lý token ngoại tuyến.", style='List Bullet')
    doc.add_paragraph("• Module Đồng bộ hóa (Synchronization Engine): Đảm bảo ghi nhận các thao tác điều chỉnh hành trình khi mất mạng và đồng bộ chớp nhoáng khi kết nối lại.", style='List Bullet')
    
    doc.add_heading("1.2 Xác định các services chính", level=2)
    doc.add_paragraph(
        "Các dịch vụ di động của ứng dụng điều phối chính các luồng truyền thông và định vị ngầm:"
    )
    doc.add_paragraph("1. TripService: Thực hiện gọi API Gateway qua luồng Server-Sent Events (SSE) để kết xuất dữ liệu thời gian thực và điều khiển cổng tái tối ưu hóa lộ trình JIT (Just-In-Time) mid-day thông qua endpoint /re_route.", style='List Number')
    doc.add_paragraph("2. BackgroundLocationService: Sử dụng thư viện expo-location và expo-task-manager đăng ký tác vụ ngầm BACKGROUND_LOCATION_TASK, lấy tọa độ thiết bị định kỳ (mỗi 60 giây hoặc 100 mét dịch chuyển).", style='List Number')
    
    doc.add_heading("1.3 Đánh giá và đề xuất hướng tổ chức state flow, xử lý lỗi/timeout", level=2)
    doc.add_paragraph(
        "Tổ chức State Flow sử dụng React Context hoặc Zustand để cô lập trạng thái đồng bộ hành trình. "
        "Để xử lý lỗi kết nối và độ trễ SSE, ứng dụng áp dụng chính sách Exponential Backoff and Jitter kết hợp "
        "kiểm tra trạng thái pre-flight qua endpoint /health trước khi nạp dữ liệu lập lịch."
    )
    
    # Mermaid block
    doc.add_paragraph("```mermaid")
    doc.add_paragraph("graph TD")
    doc.add_paragraph("    GPS[Background GPS Update] -->|Trigger 1min/100m| Check[Check Itinerary Schedule]")
    doc.add_paragraph("    Check -->|isLate == True| Push[Send Local Notification]")
    doc.add_paragraph("    Push -->|User Accepts| JIT[Call JIT Re-route Proxy]")
    doc.add_paragraph("    JIT -->|Update Polyline| Render[Render Map Route]")
    doc.add_paragraph("```")

def add_section2(doc):
    doc.add_heading("2. Nghiên cứu Khối Gateway (FastAPI Orchestrator)", level=1)
    
    doc.add_heading("2.1 Xác định các modules chính", level=2)
    doc.add_paragraph(
        "Khối Gateway đóng vai trò điều phối toàn bộ luồng nghiệp vụ của hệ thống (Layer 2 & 3), "
        "được xây dựng trên nền tảng FastAPI hoạt động bất đồng bộ tối đa. Khối này bao gồm các module chính:"
    )
    doc.add_paragraph("• Module Ingress & Bảo mật (Security & Auth Ingress): Xác thực JWT, cấu hình CORS và ngăn ngừa API abuse.", style='List Bullet')
    doc.add_paragraph("• Module Trích xuất ý định (LLM Intent Extraction): Tích hợp OpenAI với thư viện Instructor để phân tích câu lệnh tiếng Việt tự nhiên thành hợp đồng dữ liệu JSON.", style='List Bullet')
    doc.add_paragraph("• Module Bộ lọc không gian (Spatial & Semantic Hybrid Search): Truy vấn địa điểm trong PostgreSQL/PostGIS kết hợp so khớp vector ngữ nghĩa pgvector.", style='List Bullet')
    doc.add_paragraph("• Module Chấm điểm tiện ích (Utility Scorer): Chấm điểm POI dựa trên đặc tính địa điểm và hồ sơ sở thích của người dùng trước khi gửi sang bộ giải thuật lõi.", style='List Bullet')
    
    doc.add_heading("2.2 Xác định các service chính", level=2)
    doc.add_paragraph(
        "Các dịch vụ lõi của API Gateway bao gồm:"
    )
    doc.add_paragraph("1. LLMExtractorService: Dùng AsyncOpenAI + Instructor phân tích câu lệnh thành LLMDataContract. Tự động thêm nhãn 'vegetarian' nếu phát hiện yêu cầu ăn chay và loại bỏ các địa danh chung chung ra khỏi locked_pois để tránh lỗi.", style='List Number')
    doc.add_paragraph("2. SpatialFilterService & EmbeddingService: Thực hiện Hybrid Search kết hợp PostGIS ST_DWithin trên trường coordinates Geography và pgvector <=> cosine distance trên chỉ mục HNSW. Hỗ trợ mở rộng bán kính tự động FALLBACK_TIERS từ 10km đến 30km nếu số điểm không đạt MIN_POI_THRESHOLD = 10.", style='List Number')
    
    doc.add_paragraph(
        "3. UtilityScorer: Tính toán điểm hữu ích $U_i$ của địa điểm $i$ dựa trên 8 chiều trọng số độc quyền:"
    )
    
    # Mathematical equation
    doc.add_paragraph("$$U_i = 0.28 \\cdot S_{sem} + 0.18 \\cdot Q_{qual} + 0.14 \\cdot L_{loc} + 0.10 \\cdot N_{nov} + 0.10 \\cdot B_{bud} + 0.10 \\cdot C_{com} + 0.06 \\cdot D_{dist} + 0.04 \\cdot G_{div}$$")
    
    # Table of weights
    table = doc.add_table(rows=9, cols=3)
    # Style table with thin borders
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Thành phần chỉ số'
    hdr_cells[1].text = 'Ký hiệu'
    hdr_cells[2].text = 'Trọng số (%)'
    
    data = [
        ("Semantic Similarity Score", "S_sem", "28%"),
        ("Quality & Review Rating", "Q_qual", "18%"),
        ("Local Experience & Authenticity", "L_loc", "14%"),
        ("Uniqueness & Novelty Index", "N_nov", "10%"),
        ("Budget Alignment Ratio", "B_bud", "10%"),
        ("Comfort & Fatigue Rhythm", "C_com", "10%"),
        ("Geographical Proximity Score", "D_dist", "6%"),
        ("Category Diversity Marginal Gain", "G_div", "4%")
    ]
    for idx, row in enumerate(data):
        row_cells = table.rows[idx+1].cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        
    doc.add_heading("2.3 Mô tả input/output giữa Mobile và Routing Engine", level=2)
    doc.add_paragraph(
        "Ứng dụng di động gửi dữ liệu yêu cầu bằng tiếng Việt kèm GPS khách sạn và số ngày du lịch. "
        "Gateway xử lý qua luồng DB I/O Isolation (gọi LLM/Embedding không đồng bộ bên ngoài DB session, chỉ mở DB session trong 50ms để truy vấn PostGIS) "
        "rồi chuyển ma trận địa điểm tối ưu sang Layer 4 Routing Engine để chạy giải thuật OR-Tools."
    )
    
    # Sequential SSE Stream flow
    doc.add_paragraph("```mermaid")
    doc.add_paragraph("sequenceDiagram")
    doc.add_paragraph("    Mobile->>Gateway: POST /plan_trip_stream (Prompt, Hotel GPS)")
    doc.add_paragraph("    Note over Gateway: Phase A: LLM Intent Extract (No DB connection)")
    doc.add_paragraph("    Gateway-->>Mobile: SSE event: l2_done (Tags & Preferences)")
    doc.add_paragraph("    Note over Gateway: Phase B: Open DB Session for Spatial Hybrid Query (50ms)")
    doc.add_paragraph("    Gateway-->>Mobile: SSE event: l3_done (POIs Found)")
    doc.add_paragraph("    Gateway->>Routing Engine: POST /plan (Top 50 POIs)")
    doc.add_paragraph("    Routing Engine-->>Gateway: Dynamic Route coordinate chunks")
    doc.add_paragraph("    Gateway-->>Mobile: SSE event: l4_stream (Compressed Polyline chunks)")
    doc.add_paragraph("```")

def add_section3(doc):
    doc.add_heading("3. Nghiên cứu Khối Routing Engine (OR-Tools CVRP Core)", level=1)
    
    doc.add_heading("3.1 Xác định các modules chính", level=2)
    doc.add_paragraph(
        "Routing Engine đóng vai trò là Layer 4 - Bộ giải thuật lõi giải bài toán CVRPTW "
        "(Capacitated Vehicle Routing Problem with Time Windows). Các module chính bao gồm:"
    )
    doc.add_paragraph("• Module Routing API (FastAPI): Nhận yêu cầu lập lịch trình đa ngày và tái tối ưu hóa lộ trình.", style='List Bullet')
    doc.add_paragraph("• Module Phân bổ Heuristic (POI Allocator): Tiền xử lý phân chia danh sách POIs về từng ngày để giảm độ phức tạp tính toán.", style='List Bullet')
    doc.add_paragraph("• Module SQLite Distance Cache: Lưu trữ đệm khoảng cách giúp giảm 90% tần suất gọi OSRM server.", style='List Bullet')
    doc.add_paragraph("• Module OR-Tools Core Solver (ortools_impl.py): Chứa cấu trúc toán học của bài toán CVRP, cài đặt ràng buộc mệt mỏi và tránh nắng nóng ngoài trời.", style='List Bullet')
    
    doc.add_heading("3.2 Xác định các service chính", level=2)
    doc.add_paragraph(
        "Bộ giải trình lõi của Routing Engine bao gồm:"
    )
    doc.add_paragraph("1. MultiPlanner: Sinh song song 3 phương án: Balanced (Cân bằng), Budget (Tiết kiệm) và Chill (Thư thả) sử dụng cơ chế hình phạt trùng lặp (giảm 25% utility của POIs đã chọn ở phương án trước).", style='List Number')
    doc.add_paragraph("2. RestBreakInserter: Tự động chèn điểm dừng nghỉ ảo __rest_break__ 20 phút khi hành khách di chuyển/tham quan liên tục vượt quá 180 phút.", style='List Number')
    doc.add_paragraph("3. ItineraryValidator: Hậu xử lý kiểm tra lộ trình qua 4 quy tắc vàng chất lượng tour (meal_missing, consecutive_heavy, outdoor_heat, rest_needed).", style='List Number')
    
    doc.add_heading("3.3 Đề xuất routing engine tối ưu", level=2)
    doc.add_paragraph(
        "Mô hình hóa toán học của giải thuật solver được thiết lập như sau:"
    )
    doc.add_paragraph(
        "Hàm mục tiêu tối thiểu hóa tổng chi phí di chuyển kèm hình phạt nhịp điệu và fatigue, "
        "đồng thời tối đa hóa tiện ích tham quan thông qua drop penalty disjunctions tỉ lệ thuận với utility:"
    )
    doc.add_paragraph("$$\\min \\sum_{i,j,k} d_{ij} \\cdot x_{ijk} + \\sum_{k} P_{fixed} \\cdot y_k + \\sum_{\\text{diversity}} P_{div} + \\sum_{\\text{rhythm}} P_{rhy} - \\sum_{i \\in \\text{dropped}} P_{drop}(U_i)$$")
    doc.add_paragraph(
        "Trong đó $x_{ijk}$ là biến nhị phân chỉ định xe $k$ đi từ điểm $i$ sang $j$. "
        "Với các điểm không bắt buộc ghé qua, hình phạt bỏ điểm được tính theo công thức: $P_{drop}(U_i) = 100,000 + U_i \\cdot 900,000$."
    )
    
    # OR-Tools Code snippets
    doc.add_paragraph(
        "Dưới đây là mã nguồn Python thực tế của giải thuật tích hợp ràng buộc Fatigue "
        "và Tránh nắng nóng ngoài trời (12:00 - 14:00) bằng Google OR-Tools:"
    )
    
    doc.add_paragraph("```python")
    doc.add_paragraph("# 1. Ràng buộc Fatigue (Sức bền) tích lũy theo ngày")
    doc.add_paragraph("fatigue_costs = self.problem_data.get('fatigue_costs', [])")
    doc.add_paragraph("max_fatigue = self.problem_data.get('max_fatigue_per_day', 15)")
    doc.add_paragraph("if fatigue_costs:")
    doc.add_paragraph("    def fatigue_callback(from_index):")
    doc.add_paragraph("        node = manager.IndexToNode(from_index)")
    doc.add_paragraph("        if node in depot_indices: return 0")
    doc.add_paragraph("        return fatigue_costs[node] if node < len(fatigue_costs) else 0")
    doc.add_paragraph("    fatigue_cb_index = routing.RegisterUnaryTransitCallback(fatigue_callback)")
    doc.add_paragraph("    routing.AddDimensionWithVehicleCapacity(")
    doc.add_paragraph("        fatigue_cb_index, 0, [max_fatigue] * num_vehicles, True, 'Fatigue'")
    doc.add_paragraph("    )")
    doc.add_paragraph("")
    doc.add_paragraph("# 2. Ràng buộc Tránh nắng nóng trưa gắt ngoài trời (12:00 - 14:00)")
    doc.add_paragraph("avoid_outdoor = self.problem_data.get('avoid_outdoor_window')  # [720, 840] phút")
    doc.add_paragraph("is_outdoor_list = self.problem_data.get('is_outdoor_list', [])")
    doc.add_paragraph("if avoid_outdoor and is_outdoor_list:")
    doc.add_paragraph("    avoid_start_scaled = int(avoid_outdoor[0] * 100)")
    doc.add_paragraph("    avoid_end_scaled = int(avoid_outdoor[1] * 100)")
    doc.add_paragraph("    for node in range(N):")
    doc.add_paragraph("        if node in depot_indices: continue")
    doc.add_paragraph("        if node < len(is_outdoor_list) and is_outdoor_list[node]:")
    doc.add_paragraph("            index = manager.NodeToIndex(node)")
    doc.add_paragraph("            time_dimension.CumulVar(index).RemoveInterval(")
    doc.add_paragraph("                avoid_start_scaled, avoid_end_scaled")
    doc.add_paragraph("            )")
    doc.add_paragraph("```")
    
    doc.add_heading("3.4 Mô tả luồng lập kế hoạch và tái tối ưu", level=2)
    doc.add_paragraph(
        "Khi người dùng yêu cầu tái tối ưu hóa lộ trình giữa ngày do bị trễ lịch trình (kích hoạt từ Mobile GPS), "
        "hệ thống tạo một kho lưu trữ ảo (Virtual Depot) tại GPS hiện tại của người dùng, lọc bỏ các địa điểm "
        "đã đi qua và chạy lại thuật toán OR-Tools với ma trận thời gian thực mới của các POI còn lại."
    )

def add_section4_and_5(doc):
    doc.add_heading("4. Nghiên cứu Khối Data/Infrastructure", level=1)
    
    doc.add_heading("4.1 Xác định các modules chính", level=2)
    doc.add_paragraph(
        "Phân hệ Cơ sở dữ liệu và Hạ tầng đảm bảo tốc độ phản hồi và độ bền vững của hệ thống. "
        "Hạ tầng bao gồm các module chính:"
    )
    doc.add_paragraph("• Module PostgreSQL + PostGIS (Spatial Data): Quản lý lưu trữ không gian của các địa điểm du lịch.", style='List Bullet')
    doc.add_paragraph("• Module pgvector (Semantic Search): Tìm kiếm tương đồng vector bằng chỉ mục HNSW.", style='List Bullet')
    doc.add_paragraph("• Module OSRM (Open Source Routing Machine): Tự cài đặt server cục bộ cho dữ liệu bản đồ Thừa Thiên Huế để tính toán ma trận khoảng cách siêu nhanh.", style='List Bullet')
    
    doc.add_heading("4.2 Xác định các service chính", level=2)
    doc.add_paragraph("1. AsyncSessionFactory (SQLAlchemy): Cơ chế kết nối không đồng bộ tới DB PostgreSQL, đảm bảo đóng kết nối ngay lập tức sau 50ms truy vấn không gian để tránh nghẽn Connection Pool.", style='List Number')
    doc.add_paragraph("2. Alembic: Quản lý di cư phiên bản DB tự động và lưu trữ trường vector nhúng 1536 chiều.", style='List Number')
    doc.add_paragraph("3. OSRM Service: Chạy trong Docker container nạp tệp hue.osrm, phục vụ API tính toán khoảng cách/travel-time ma trận nhanh hơn 20 lần so với các dịch vụ đám mây ngoại vi.", style='List Number')
    
    doc.add_heading("4.3 Đề xuất hạ tầng chạy local phù hợp", level=2)
    doc.add_paragraph(
        "Toàn bộ hạ tầng hệ thống được đóng gói bằng Docker-Compose để chạy local ổn định:"
    )
    doc.add_paragraph("• PostgreSQL + PostGIS + pgvector: Cổng 5432.", style='List Bullet')
    doc.add_paragraph("• OSRM Backend (hue.osrm): Cổng 5001.", style='List Bullet')
    doc.add_paragraph("• Core Routing Engine (FastAPI): Cổng 8000.", style='List Bullet')
    doc.add_paragraph("• API Gateway (FastAPI): Cổng 8001.", style='List Bullet')
    doc.add_paragraph("• Web Dashboard UI (Next.js): Cổng 3000.", style='List Bullet')
    
    doc.add_heading("4.4 Mô tả dependency hạ tầng và luồng dữ liệu", level=2)
    doc.add_paragraph(
        "Luồng trao đổi dữ liệu toàn cục bắt đầu từ người dùng di động -> API Gateway (tách DB session gọi LLM & Embedding) "
        "-> truy vấn PostGIS/pgvector lấy danh sách POI tối ưu -> truyền Top 50 điểm sang Routing Engine -> "
        "Routing Engine sử dụng SQLite Cache / OSRM để dựng ma trận chi phí -> gọi OR-Tools giải thuật tìm lộ trình tốt nhất -> "
        "trả kết quả SSE polylines về Mobile App."
    )
    
    doc.add_heading("5. Tổng hợp và kết luận", level=1)
    
    doc.add_heading("5.1 Tóm tắt theo khối", level=2)
    doc.add_paragraph("• Mobile: Đã hoàn thành cấu trúc định vị ngầm và kết nối SSE luồng dữ liệu tối ưu thời gian thực.", style='List Bullet')
    doc.add_paragraph("• Gateway: Hoàn thành bộ trích xuất ý định LLM và cơ chế tìm kiếm lai Hybrid Search PostGIS/pgvector.", style='List Bullet')
    doc.add_paragraph("• Routing Engine: Hoàn thành cấu hình CVRPTW với các ràng buộc sức bền, nắng nóng ngoài trời gắt và tự động chèn điểm nghỉ ngơi.", style='List Bullet')
    doc.add_paragraph("• Data/Infrastructure: Đóng gói thành công toàn bộ hạ tầng cơ sở dữ liệu và OSRM tự host qua Docker cục bộ.", style='List Bullet')
    
    doc.add_heading("5.2 Danh sách vấn đề cần xử lý tiếp", level=2)
    doc.add_paragraph("1. Độ phủ bản đồ OSRM: Dữ liệu giao thông Huế cần được cập nhật thường xuyên từ bản đồ OpenStreetMap để tránh lỗi định vị đường cụt.", style='List Number')
    doc.add_paragraph("2. Tối ưu hóa pin: Theo dõi định vị ngầm trên Mobile cần được thiết lập ở mức Balanced để bảo vệ thời lượng pin của khách du lịch.", style='List Number')
    
    doc.add_heading("5.3 Đề xuất bước tiếp theo", level=2)
    doc.add_paragraph("• Triển khai bộ kiểm thử tự động E2E toàn bộ luồng tích hợp hệ thống từ Mobile đến Core Solver.", style='List Bullet')
    doc.add_paragraph("• Mở rộng tập dữ liệu POIs sang các thành phố du lịch lân cận như Đà Nẵng, Hội An để kiểm nghiệm hiệu suất giải thuật.", style='List Bullet')

def create_document():
    doc = Document()
    
    # Configure margins (Left=3.0cm, Right=2.0cm, Top=2.0cm, Bottom=2.0cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7874)    # 2.0 cm
        section.bottom_margin = Inches(0.7874) # 2.0 cm
        section.left_margin = Inches(1.1811)   # 3.0 cm
        section.right_margin = Inches(0.7874)  # 2.0 cm
        
    # Set default Normal style font
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Create or update custom Heading Styles
    def configure_heading(name, size, space_before, space_after, bold=True, italic=False):
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(size)
        font.bold = bold
        font.italic = italic
        font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.line_spacing = 1.15
        return style
        
    configure_heading('Heading 1', 16, 12, 6)
    configure_heading('Heading 2', 13, 12, 4)
    configure_heading('Heading 3', 12, 6, 2, bold=True, italic=True)
    
    # Title Style
    configure_heading('Title', 24, 18, 12)
    
    # Build sections
    add_title_and_section1(doc)
    add_section2(doc)
    add_section3(doc)
    add_section4_and_5(doc)
    
    return doc

if __name__ == "__main__":
    os.makedirs("scripts", exist_ok=True)
    doc = create_document()
    doc.save("AI_Travel_Optimizer_Architecture_Report.docx")
    print("Report generated successfully!")
